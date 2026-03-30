"""
Utility functions for MCQ and question generation
"""

import re
import io
from fastapi import HTTPException
from docx import Document
from typing import List, Dict, Optional
from pypdf import PdfReader
import os 
import shutil
from src.config import LLM_CONFIG, MAX_DOCUMENT_LENGTH, ALLOWED_FILE_TYPES, SCRIPT_MODIFICATION_PROMPT_TEMPLATE

from docx import Document
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser, StrOutputParser
from langchain_core.pydantic_v1 import BaseModel, Field
from gtts import gTTS
from PIL import Image, ImageDraw, ImageFont
from moviepy import ImageClip, AudioFileClip, concatenate_videoclips
import zipfile
from groq import Groq

# ============================================================================
# DOCUMENT UTILITIES
# ============================================================================

def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from a Word document.
    
    Args:
        content: Raw bytes of the .docx file
        
    Returns:
        Extracted text as string
        
    Raises:
        HTTPException: If document cannot be read or is empty
    """
    try:
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        # Extract from paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
                
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        document_text = "\n\n".join(text_parts)
        
        if not document_text:
            raise ValueError("Document contains no readable text")
        
        return document_text
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX document: {str(e)}")


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from a PDF document using PyPDF.
    
    Args:
        content: Raw bytes of the .pdf file
        
    Returns:
        Extracted text as string
        
    Raises:
        HTTPException: If document cannot be read or is empty
    """
    try:
        text_parts = []
        
        pdf_reader = PdfReader(io.BytesIO(content))
        
        if len(pdf_reader.pages) == 0:
            raise HTTPException(status_code=400, detail="PDF document is empty")
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
        
        document_text = "\n\n".join(text_parts)
        
        if not document_text:
            raise HTTPException(status_code=400, detail="No extractable text found in PDF")
        
        return document_text
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read PDF document: {str(e)}")


def extract_text_from_file(filename: str, content: bytes) -> str:
    """
    Extract text from document based on file type.
    
    Args:
        filename: Name of the file
        content: Raw bytes of the file
        
    Returns:
        Extracted text as string
        
    Raises:
        HTTPException: If file type is not supported or cannot be read
    """
    file_ext = "." + filename.split('.')[-1].lower()
    
    if file_ext == ".docx":
        return extract_text_from_docx(content)
    elif file_ext == ".pdf":
        return extract_text_from_pdf(content)
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file_ext}. Supported types: .docx, .pdf"
        )





# ============================================================================
# VALIDATION UTILITIES
# ============================================================================

def validate_file(filename: str, allowed_types: List[str]) -> None:
    """
    Validate uploaded file type.
    
    Args:
        filename: Name of the uploaded file
        allowed_types: List of allowed file extensions (e.g., ['.docx'])
        
    Raises:
        HTTPException: If file type is not allowed
    """
    file_ext = "." + filename.split('.')[-1].lower()
    if file_ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"File must be one of {allowed_types}. Got: {file_ext}"
        )


def validate_num_questions(num_questions: int, min_val: int, max_val: int) -> None:
    """
    Validate number of questions parameter.
    
    Args:
        num_questions: Number of questions to generate
        min_val: Minimum allowed questions
        max_val: Maximum allowed questions
        
    Raises:
        HTTPException: If num_questions is out of valid range
    """
    if num_questions < min_val or num_questions > max_val:
        raise HTTPException(
            status_code=400,
            detail=f"num_questions must be between {min_val} and {max_val}"
        )


def validate_document_length(text: str, min_length: int) -> None:
    """
    Validate minimum document length.
    
    Args:
        text: Document text
        min_length: Minimum required character count
        
    Raises:
        HTTPException: If document is too short
    """
    if len(text) < min_length:
        raise HTTPException(
            status_code=400,
            detail=f"Document too short. Minimum {min_length} characters required."
        )


# ============================================================================
# PARSING UTILITIES
# ============================================================================

def parse_mcqs(text: str) -> List[Dict]:
    """
    Parse MCQ text response into structured format.
    
    Args:
        text: Raw LLM response containing MCQs
        
    Returns:
        List of dictionaries with question, options, and correct answer
    """
    mcqs = []
    
    # Split by question pattern
    questions = re.split(r'Question \d+:', text)
    
    for q in questions[1:]:  # Skip first empty split
        lines = q.strip().split('\n')
        if len(lines) < 5:
            continue
        
        question_text = lines[0].strip()
        options = {}
        correct_answer = None
        
        # Parse options and answer
        for line in lines[1:]:
            line = line.strip()
            if line.startswith('A)'):
                options['A'] = line[2:].strip()
            elif line.startswith('B)'):
                options['B'] = line[2:].strip()
            elif line.startswith('C)'):
                options['C'] = line[2:].strip()
            elif line.startswith('D)'):
                options['D'] = line[2:].strip()
            elif line.startswith('Correct Answer:'):
                correct_answer = line.split(':')[1].strip()
        
        # Only add if all required fields are present
        if question_text and len(options) == 4 and correct_answer:
            mcqs.append({
                "question": question_text,
                "options": options,
                "correct_answer": correct_answer
            })
    
    return mcqs


def parse_questions(text: str) -> List[Dict]:
    """
    Parse general questions from LLM response.
    
    Args:
        text: Raw LLM response containing questions
        
    Returns:
        List of dictionaries with questions and answers
    """
    questions = []
    
    # Split by question pattern
    q_splits = re.split(r'Question \d+:', text)
    
    for q in q_splits[1:]:  # Skip first empty split
        lines = q.strip().split('\n')
        if len(lines) < 2:
            continue
        
        question_text = lines[0].strip()
        answer_text = None
        
        # Find answer line
        for i, line in enumerate(lines[1:], 1):
            if line.strip().lower().startswith('answer:'):
                answer_text = line.split(':', 1)[1].strip()
                break
        
        if question_text and answer_text:
            questions.append({
                "question": question_text,
                "answer": answer_text
            })
    
    return questions


def parse_fill_in_the_blanks(text: str) -> List[Dict]:
    """
    Parse fill-in-the-blanks questions from LLM response.
    
    Args:
        text: Raw LLM response containing fill-in-the-blanks questions
        
    Returns:
        List of dictionaries with questions, answers, and context
    """
    blanks = []
    
    # Split by question pattern
    q_splits = re.split(r'Question \d+:', text)
    
    for q in q_splits[1:]:  # Skip first empty split
        lines = q.strip().split('\n')
        if len(lines) < 2:
            continue
        
        question_text = None
        blank_answer = None
        context = None
        
        # Parse the lines
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            if question_text is None:
                # First non-empty line is the question
                question_text = line_stripped
            elif line_stripped.lower().startswith('blank answer:'):
                blank_answer = line_stripped.split(':', 1)[1].strip()
            elif line_stripped.lower().startswith('context:'):
                context = line_stripped.split(':', 1)[1].strip()
        
        # Only add if required fields are present
        if question_text and blank_answer:
            blanks.append({
                "question": question_text,
                "blank_answer": blank_answer,
                "context": context
            })
    
    return blanks


# ============================================================================
# TEXT UTILITIES
# ============================================================================

def truncate_text(text: str, max_length: int) -> str:
    """
    Truncate text to maximum length.
    
    Args:
        text: Input text
        max_length: Maximum length
        
    Returns:
        Truncated text
    """
    if len(text) <= max_length:
        return text
    return text[:max_length] + "..."


def clean_text(text: str) -> str:
    """
    Clean text by removing extra whitespace.
    
    Args:
        text: Input text
        
    Returns:
        Cleaned text
    """
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def cleanup_files(filepath: str, dirpath: str):
    """Delete the generated video file and temporary directory."""
    import gc
    import time
    # Force GC to release any file handles held by objects
    gc.collect()
    
    try:
        if filepath and os.path.exists(filepath):
            # Try a couple of times for Windows locks
            for _ in range(3):
                try:
                    os.remove(filepath)
                    break
                except Exception:
                    time.sleep(0.5)
                
        if dirpath and os.path.exists(dirpath):
            # Try a couple of times for Windows locks
            for _ in range(3):
                shutil.rmtree(dirpath, ignore_errors=True)
                if not os.path.exists(dirpath):
                    break
                time.sleep(0.5)
    except Exception as e:
        print(f"Error during cleanup: {e}")



def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from a Word document.
    
    Args:
        content: Raw bytes of the .docx file
        
    Returns:
        Extracted text as string
        
    Raises:
        HTTPException: If document cannot be read or is empty
    """
    try:
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        # Extract from paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
                
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        document_text = "\n\n".join(text_parts)
        
        if not document_text:
            raise ValueError("Document contains no readable text")
        
        return document_text
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read document: {str(e)}")


def validate_file(filename: str) -> None:
    """
    Validate uploaded file type.
    
    Args:
        filename: Name of the uploaded file
        
    Raises:
        HTTPException: If file type is not allowed
    """
    file_ext = "." + filename.split('.')[-1].lower()
    if file_ext not in ALLOWED_FILE_TYPES:
        raise HTTPException(
            status_code=400, 
            detail=f"File must be a .docx or .pdf file. Got: {file_ext}"
        )


def validate_num_questions(num_questions: int) -> None:
    """
    Validate number of questions parameter.
    
    Args:
        num_questions: Number of questions to generate
        
    Raises:
        HTTPException: If num_questions is out of valid range
    """
    if num_questions < MIN_QUESTIONS or num_questions > MAX_QUESTIONS:
        raise HTTPException(
            status_code=400, 
            detail=f"num_questions must be between {MIN_QUESTIONS} and {MAX_QUESTIONS}"
        )


def initialize_llm(groq_api_key: str):
    """
    Initialize and return Groq LLM instance.
    
    Args:
        groq_api_key: API key for Groq service
        
    Returns:
        Initialized ChatGroq instance
    """
    return ChatGroq(
        model=LLM_CONFIG["model"],
        temperature=LLM_CONFIG["temperature"],
        max_tokens=LLM_CONFIG["max_tokens"],
        groq_api_key=groq_api_key
    )


def create_prompt_chain(llm):
    """
    Create LangChain prompt and execution chain.
    
    Args:
        llm: Initialized ChatGroq instance
        
    Returns:
        Execution chain combining prompt template and LLM
    """
    prompt = PromptTemplate(
        input_variables=["document_text", "num_questions"],
        template=MCQ_PROMPT_TEMPLATE
    )
    return prompt | llm | StrOutputParser()


def draw_rounded_rectangle(draw, xy, corner_radius, fill=None, outline=None):
    """Draw a rounded rectangle."""
    upper_left_point = xy[0]
    bottom_right_point = xy[1]
    draw.pieslice([upper_left_point, (upper_left_point[0] + corner_radius * 2, upper_left_point[1] + corner_radius * 2)],
                  180, 270, fill=fill, outline=outline)
    draw.pieslice([(bottom_right_point[0] - corner_radius * 2, upper_left_point[1]), (bottom_right_point[0], upper_left_point[1] + corner_radius * 2)],
                  270, 360, fill=fill, outline=outline)
    draw.pieslice([(upper_left_point[0], bottom_right_point[1] - corner_radius * 2), (upper_left_point[0] + corner_radius * 2, bottom_right_point[1])],
                  90, 180, fill=fill, outline=outline)
    draw.pieslice([(bottom_right_point[0] - corner_radius * 2, bottom_right_point[1] - corner_radius * 2), (bottom_right_point[0], bottom_right_point[1])],
                  0, 90, fill=fill, outline=outline)
    draw.rectangle([(upper_left_point[0], upper_left_point[1] + corner_radius), (bottom_right_point[0], bottom_right_point[1] - corner_radius)],
                   fill=fill, outline=outline)
    draw.rectangle([(upper_left_point[0] + corner_radius, upper_left_point[1]), (bottom_right_point[0] - corner_radius, bottom_right_point[1])],
                   fill=fill, outline=outline)

def create_slide_image(title, bullets, output_path, logo_path=None):
    """Create a slide image using Pillow with custom green wave aesthetic."""
    import math
    import textwrap
    import os
    
    width, height = 1280, 720
    
    # White background for contrast
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    overlay = Image.new('RGBA', (width, height), (0,0,0,0))
    draw = ImageDraw.Draw(overlay)
    
    c_green = (85, 175, 90)
    
    # Top Banner - Height 160
    draw.rectangle([0, 0, width, 160], fill=c_green)
    
    # Bottom Banner - Height 55
    draw.rectangle([0, height-55, width, height], fill=c_green)
    
    # Add Logo White Box (Rounded Rectangle) - Increased Size
    box_w, box_h = 150, 150
    box_x = width - box_w - 40
    box_y = 10
    draw_rounded_rectangle(draw, [(box_x, box_y), (box_x + box_w, box_y + box_h)], corner_radius=10, fill=(255, 255, 255, 255))

    img.paste(overlay, mask=overlay)
    draw = ImageDraw.Draw(img)

    if logo_path and os.path.exists(logo_path):
        try:
            logo = Image.open(logo_path).convert("RGBA")
            # Resize logo to fit inside the white box with padding
            max_size = 135
            logo.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
            
            logo_x = box_x + (box_w - logo.width) // 2
            logo_y = box_y + (box_h - logo.height) // 2
            
            # Use 'logo' as mask if it has an alpha channel, else paste normally
            img.paste(logo, (logo_x, logo_y), mask=logo if logo.mode == 'RGBA' else None)
        except Exception as e:
            print(f"Error drawing logo: {e}")

    # Use bold for title
    try:
        # Increased title size to 42 for better presence as in sample_template.png
        title_font = ImageFont.truetype("arialbd.ttf", 42)
        content_font = ImageFont.truetype("arial.ttf", 28)
    except:
        title_font = ImageFont.load_default()
        content_font = ImageFont.load_default()
        
    # Text colors
    title_color = (255, 255, 255)
    text_color = (0, 0, 0)
    
    # Title wrapping - slightly wider for larger font
    title_wrapped = textwrap.wrap(title, width=32)
    y_offset = 45 # Adjusted for 160px header
    if len(title_wrapped) == 1:
        y_offset = 55
    
    for line in title_wrapped:
        draw.text((60, y_offset), line, font=title_font, fill=title_color)
        y_offset += 50
        
    # --- Draw Bullets ---
    y_offset = 240
    safe_bullet_chars = int((width - 140) / 15)  # Adjusted for smaller font
    for bullet in bullets:
        if y_offset > height - 120:  # Restrict to avoid footer overlap
            break
        bullet_wrapped = textwrap.wrap(f"\u2022 {bullet}", width=safe_bullet_chars)
        for i, line in enumerate(bullet_wrapped):
            if y_offset > height - 120:
                break
            indent = 0 if i == 0 else 20
            draw.text((100 + indent, y_offset), line, font=content_font, fill=text_color)
            y_offset += 40  # Less vertical spacing for smaller font
        y_offset += 20

    img.save(output_path)






def parse_mcqs(text: str) -> list:
    """
    Parse MCQ text response into structured format.
    
    Args:
        text: Raw LLM response containing MCQs
        
    Returns:
        List of dictionaries with question, options, and correct answer
    """
    mcqs = []
    
    # Split by question pattern
    questions = re.split(r'Question \d+:', text)
    
    for q in questions[1:]:  # Skip first empty split
        lines = q.strip().split('\n')
        if len(lines) < 5:
            continue
        
        question_text = lines[0].strip()
        options = {}
        correct_answer = None
        
        # Parse options and answer
        for line in lines[1:]:
            line = line.strip()
            if line.startswith('A)'):
                options['A'] = line[2:].strip()
            elif line.startswith('B)'):
                options['B'] = line[2:].strip()
            elif line.startswith('C)'):
                options['C'] = line[2:].strip()
            elif line.startswith('D)'):
                options['D'] = line[2:].strip()
            elif line.startswith('Correct Answer:'):
                correct_answer = line.split(':')[1].strip()
        
        # Only add if all required fields are present
        if question_text and len(options) == 4 and correct_answer:
            mcqs.append({
                "question": question_text,
                "options": options,
                "correct_answer": correct_answer
            })
    
    return mcqs




def create_script_txt(slides: List[Dict], output_path: str):
    """Create a plain text file containing the slide titles and scripts."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("VIDEO SCRIPT\n")
        f.write("=" * 20 + "\n\n")
        
        for i, slide in enumerate(slides):
            f.write(f"SLIDE {i+1}: {slide['title']}\n")
            f.write("-" * (8 + len(str(i+1)) + len(slide['title'])) + "\n")
            f.write(f"NARRATIVE: {slide['script']}\n\n")
            
            if slide.get('bullets'):
                f.write("KEY POINTS:\n")
                for bullet in slide['bullets']:
                    f.write(f"  - {bullet}\n")
                f.write("\n")
            

            
            f.write("\n")

def parse_script_txt(content: str) -> List[Dict]:
    """Parse the edited slides_content.txt back into a list of slide dictionaries."""
    slides = []
    # Split content into slides using "SLIDE X:" as delimiter
    # Use regex to find "SLIDE " followed by a number and a colon
    slide_blocks = re.split(r'SLIDE \d+:', content)
    
    # slide_blocks[0] will contain the header "VIDEO SCRIPT..."
    for block in slide_blocks[1:]:
        lines = block.strip().split('\n')
        if not lines:
            continue
            
        # First line is the title (as the "SLIDE X: " part was removed by split)
        title = lines[0].strip()
        script = ""
        bullets = []
        
        current_section = None
        for line in lines[1:]:
            line = line.strip()
            # Skip empty lines and separators
            if not line or line.startswith('-' * 5):
                continue
            
            if line.startswith("NARRATIVE:"):
                script = line.replace("NARRATIVE:", "").strip()
                current_section = "narrative"
            elif line.startswith("KEY POINTS:"):
                current_section = "bullets"
            elif line.startswith("-") and current_section == "bullets":
                # Remove the leading "-" and optional space
                bullets.append(line[1:].strip())
            elif current_section == "narrative":
                # Handle multi-line narrative
                script += " " + line
                
        if title:
            slide_dict = {
                "title": title,
                "script": script.strip(),
                "bullets": bullets
            }
            slides.append(slide_dict)
    return slides


def transcribe_audio(audio_path: str, groq_api_key: str) -> str:
    """
    Transcribe audio file using Groq Whisper model.
    Returns plain text transcription.
    """
    client = Groq(api_key=groq_api_key)
    try:
        with open(audio_path, "rb") as file:
            transcription = client.audio.transcriptions.create(
                file=(audio_path, file.read()),
                model="whisper-large-v3",
                response_format="text",
            )
            return transcription
    except Exception as e:
        print(f"Transcription error for {audio_path}: {e}")
        return ""


def transcribe_audio_with_timestamps(audio_path: str, groq_api_key: str) -> list:
    """
    Transcribe audio using Groq Whisper and return segment-level timestamps.
    Returns a list of dicts: [{start, end, text}, ...]
    """
    client = Groq(api_key=groq_api_key)
    try:
        with open(audio_path, "rb") as file:
            result = client.audio.transcriptions.create(
                file=(audio_path, file.read()),
                model="whisper-large-v3",
                response_format="verbose_json",
                timestamp_granularities=["segment"]
            )
        segments = []
        if hasattr(result, 'segments') and result.segments:
            for seg in result.segments:
                # Groq SDK may return segments as dicts or as objects
                def _get(s, key):
                    return s[key] if isinstance(s, dict) else getattr(s, key)
                segments.append({
                    "start": _get(seg, "start"),
                    "end": _get(seg, "end"),
                    "text": _get(seg, "text").strip()
                })
        return segments
    except Exception as e:
        print(f"Timestamped transcription error for {audio_path}: {e}")
        return []


def create_subtitle_clip(text: str, start: float, duration: float, width: int = 1280, height: int = 720):
    """
    Create a MoviePy ImageClip that renders a subtitle bar at the bottom of the frame.
    Returns an ImageClip positioned at the bottom of the slide.
    """
    import numpy as np
    import textwrap
    from moviepy import ImageClip as MpImageClip

    bar_height = 80
    padding = 12

    # Create full-size transparent RGBA image
    img = Image.new("RGBA", (width, bar_height), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    # Semi-transparent black bar background
    # Removed the background bar to make subtitles display directly on the slide
    # draw.rectangle([0, 0, width, bar_height], fill=(0, 0, 0, 180))

    # Font
    try:
        font = ImageFont.truetype("arial.ttf", 28)
    except Exception:
        font = ImageFont.load_default()

    # Wrap text
    wrapped = textwrap.fill(text, width=90)
    lines = wrapped.split("\n")

    # Draw each line centered
    line_h = 32
    total_h = len(lines) * line_h
    y = (bar_height - total_h) // 2
    for line in lines:
        try:
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
        except Exception:
            tw = len(line) * 15
        x = (width - tw) // 2
        # Draw white outline for readability of black text on potentially busy/colored backgrounds
        draw.text((x + 2, y + 2), line, font=font, fill=(255, 255, 255, 255))
        draw.text((x - 1, y - 1), line, font=font, fill=(255, 255, 255, 255))
        draw.text((x + 1, y - 1), line, font=font, fill=(255, 255, 255, 255))
        draw.text((x - 1, y + 1), line, font=font, fill=(255, 255, 255, 255))
        draw.text((x + 1, y + 1), line, font=font, fill=(255, 255, 255, 255))
        # Main text in SOLID BLACK as requested
        draw.text((x, y), line, font=font, fill=(0, 0, 0, 255))
        y += line_h

    # Convert RGBA PIL directly to a numpy array
    # MoviePy's ImageClip will automatically read the alpha channel as a mask
    frame = np.array(img)

    bottom_margin = 80  # lift subtitles just above the bottom edge but avoid overlapping slide content
    clip = MpImageClip(frame).with_duration(duration).with_start(start).with_position((0, height - bar_height - bottom_margin))
    return clip

def parse_transcription_txt(content: str) -> Dict[int, str]:
    """
    Parse the transcription.txt file into a dictionary mapping slide index to text.
    Expected format:
    Slide 1:
    [text]
    
    Slide 2:
    [text]
    """
    transcription_map = {}
    # Split by "Slide X:" pattern
    blocks = re.split(r'Slide (\d+):', content)
    
    # blocks[0] is everything before the first "Slide X:"
    for i in range(1, len(blocks), 2):
        slide_num = int(blocks[i])
        text = blocks[i+1].strip()
        transcription_map[slide_num] = text
        
    return transcription_map

def merge_transcription_with_slides(slides: List[Dict], transcription: Dict[int, str]) -> List[Dict]:
    """
    Update the 'script' field of each slide with text from the transcription map.
    """
    updated_slides = []
    for i, slide in enumerate(slides):
        slide_num = i + 1
        new_script = transcription.get(slide_num, slide.get('script', ''))
        
        updated_slide = slide.copy()
        updated_slide['script'] = new_script
        updated_slides.append(updated_slide)
        
    return updated_slides

def modify_script_with_llm(script_text: str, user_prompt: str, groq_api_key: str, source_document: str = "") -> tuple:
    """
    Use LLM to modify the video script based on a user prompt and optional source document context.
    Returns (modified_script, modification_summary).
    """
    llm = initialize_llm(groq_api_key)
    prompt_template = PromptTemplate(
        template=SCRIPT_MODIFICATION_PROMPT_TEMPLATE,
        input_variables=["original_script", "user_prompt", "source_document"]
    )
    chain = prompt_template | llm | StrOutputParser()
    
    try:
        raw_response = chain.invoke({
            "original_script": script_text,
            "user_prompt": user_prompt,
            "source_document": source_document if source_document else "No source document provided."
        })
        
        raw_response = raw_response.strip()
        summary = "No summary provided."
        modified_script = raw_response
        
        if "[SUMMARY]" in raw_response and "[SCRIPT]" in raw_response:
            try:
                parts = raw_response.split("[SCRIPT]")
                summary = parts[0].replace("[SUMMARY]", "").strip()
                modified_script = parts[1].strip()
            except Exception as e:
                print(f"Error parsing LLM response parts: {e}")
        
        return modified_script, summary
    except Exception as e:
        print(f"Error modifying script with LLM: {e}")
        raise RuntimeError(f"Failed to modify script: {e}")


